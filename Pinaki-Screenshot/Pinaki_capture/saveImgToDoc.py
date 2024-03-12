from docx import Document
from docx.shared import Inches
import pyautogui
import os
from pynput.keyboard import Key, Listener
import datetime

currentDir = os.getcwd() 
document = Document()

screenshotShotcut= Key.print_screen  
exit_combination_msg='ctrl_l(left ctrl) + Esc' 
exit_combination = {Key.ctrl_l, Key.esc} 
defaultDocName = "testingScreenShots"
imgCount = 1 



currently_pressed = set()


def on_press(key):
	check_key(key)
	if key in exit_combination:
		currently_pressed.add(key)
		if currently_pressed == exit_combination:
			listener.stop()
			exitFun()
			
		
def on_release(key):	
	try:
		currently_pressed.remove(key)
	except KeyError:
		pass 
		
def check_key(key):
	if key == screenshotShotcut:
		saveImg() 


def addSSToDoc(imgPath):
	p = document.add_paragraph()
	r = p.add_run() 
	r.add_picture(imgPath,width=Inches(5.90551), height=Inches(3.54331)) 
	p = document.add_paragraph() 

def saveDoc(fileName):
	document.save(MasterPath +"\\"+fileName+'.docx') 
	print()
	print("Document Saved at =>")
	print(MasterPath +"\\"+fileName+".docx ") 

def exitFun():
	print("<===================================>")
	print("Press Return Key to save Document name as Default name ("+defaultDocName+")")
	fileName = input("Enter Document Name : ") 
	bad_chars = [';', ':', '!', "*","<",">","\"","/","|","?","*"] 
	for i in bad_chars : 
		fileName = fileName.replace(i, '')

	if(fileName==''): 
		fileName = defaultDocName
	saveDoc(fileName)


def createDirs():
	global MasterPath
	now = datetime.datetime.now()
	currentDateTime = now.strftime("%d.%m.%y_%I%M%S%p")
	MasterPath = currentDir+"\\TestingData\\"+currentDateTime 
	os.makedirs(MasterPath+"\\"+ "shots") 

def saveImg():
	global imgCount
	shot = pyautogui.screenshot() 
	path = MasterPath +"\\shots\\" 
	shot.save(path+'\\'+str(imgCount)+'.png') 
	addSSToDoc(path+'//'+str(imgCount)+'.png') 
	print('File Saved as ' +path+str(imgCount)+'.png') 
	imgCount=imgCount +1 


def printStartMsg():
	print("<---------------------------------------------------->")
	print(" Created by Pinaki Shashishekhar Mathan ")
	print("_______________________")
	print("Press PrtSc to take the Screenshot and save to folder and To Document")
	print("Press "+exit_combination_msg +" to exit and save the document")
	print("_______________________")
	print("Current Directory is =>"+os.getcwd())
	createDirs()
	

with Listener(on_press=on_press,on_release=on_release) as listener:
	printStartMsg()	
	listener.join()

